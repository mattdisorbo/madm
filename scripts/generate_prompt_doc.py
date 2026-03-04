"""Generate a Word doc showing all 16 prompts (8 datasets x base/adversarial)."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

doc.add_heading('MADM Prompt Reference', level=0)
doc.add_paragraph(
    'All 16 prompts used in Stage 0 (8 datasets x base/adversarial). '
    'Both paths use reasoning in step 1. The ONLY difference is that adversarial '
    'adds a critique step ("Critique this reasoning in 1 sentence") before escalation.'
)

datasets = [
    ("FEVEROUS", {
        "base_prompt": "You are helping Person 1 with fact-checking. This statement needs to be checked: {claim}.",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (support) or 'PREDICTION: 0' (refute)."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{pred_response}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is considered correct if it equals Person 1's true decision."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (support) or 'PREDICTION: 0' (refute)."),
            ("Critique",
             "SCENARIO:\n{base_prompt}\n\nPROPOSED LOGIC:\n{thought}\n\n"
             "TASK: Critique this reasoning in 1 sentence."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{thought}\n\nCRITIQUE:\n{critique}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is considered correct if it equals Person 1's true decision."),
        ],
    }),
    ("LendingClub", {
        "base_prompt": "You are helping Person 1, a loan officer, decide whether to approve a loan application. "
                       "The applicant is requesting ${loan_amnt} for {purpose}. They have been employed for {emp_length}, "
                       "a debt-to-income ratio of {dti}%, and a credit score of {fico}.",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (approve) or 'PREDICTION: 0' (reject)."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{pred_response}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is considered correct if it equals Person 1's true decision."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (approve) or 'PREDICTION: 0' (reject)."),
            ("Critique",
             "SCENARIO:\n{base_prompt}\n\nPROPOSED LOGIC:\n{thought}\n\n"
             "TASK: Critique this reasoning in 1 sentence."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{thought}\n\nCRITIQUE:\n{critique}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is considered correct if it equals Person 1's true decision."),
        ],
    }),
    ("MoralMachine", {
        "base_prompt": "An autonomous vehicle is about to get in an accident. "
                       "If the car {action_0}, {n_chars_0} {attr_0} will die. "
                       "If the car {action_1}, {n_chars_1} {attr_1} will die. "
                       "[conditional: pedestrian crossing signal info] "
                       "Person 1, with the following characteristics, is in the driver's seat: "
                       "Age: {age}. Education: {education}. Gender: {gender}. Income: {income}. "
                       "Political (0 is Conservative, 1 is Progressive): {political}. "
                       "Religious (0 is Not Religious, 1 is Religious): {religious}.",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (swerve) or 'PREDICTION: 0' (do not swerve)."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{pred_response}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is considered correct if it equals Person 1's true decision."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (swerve) or 'PREDICTION: 0' (do not swerve)."),
            ("Critique",
             "SCENARIO:\n{base_prompt}\n\nPROPOSED LOGIC:\n{thought}\n\n"
             "TASK: Critique this reasoning in 1 sentence."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{thought}\n\nCRITIQUE:\n{critique}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is considered correct if it equals Person 1's true decision."),
        ],
    }),
    ("MovieLens", {
        "base_prompt": "Person 1 has reviewed the following movies:\n"
                       "[rating history]\n"
                       "Consider these two movies they have not seen:\n"
                       "[movie pair with genres]",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict which movie Person 1 would rate higher.\n"
             "Explain your reasoning in 1 sentence. Then conclude with exactly:\n"
             "'PREDICTION: 1' (first movie) or 'PREDICTION: 2' (second movie)."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{pred_response}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is correct if it matches the movie Person 1 would rate higher."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict which movie Person 1 would rate higher.\n"
             "Explain your reasoning in 1 sentence. Then conclude with exactly:\n"
             "'PREDICTION: 1' (first movie) or 'PREDICTION: 2' (second movie)."),
            ("Critique",
             "SCENARIO:\n{base_prompt}\n\nPROPOSED LOGIC:\n{thought}\n\n"
             "TASK: Critique this reasoning in 1 sentence."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{thought}\n\nCRITIQUE:\n{critique}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is correct if it matches the movie Person 1 would rate higher."),
        ],
    }),
    ("WikipediaToxicity", {
        "base_prompt": "You are helping a group of crowd-workers label Wikipedia discussion comments as toxic or not. "
                       "This comment needs to be checked: '{comment}'.",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict the crowd-workers' aggregate decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (toxic) or 'PREDICTION: 0' (not toxic)."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{pred_response}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to the crowd-workers.\n"
             "A decision is considered correct if it equals the crowd-workers' true aggregate decision."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict the crowd-workers' aggregate decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (toxic) or 'PREDICTION: 0' (not toxic)."),
            ("Critique",
             "SCENARIO:\n{base_prompt}\n\nPROPOSED LOGIC:\n{thought}\n\n"
             "TASK: Critique this reasoning in 1 sentence."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{thought}\n\nCRITIQUE:\n{critique}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to the crowd-workers.\n"
             "A decision is considered correct if it equals the crowd-workers' true aggregate decision."),
        ],
    }),
    ("HotelBookings", {
        "base_prompt": "You are helping predict Person 1's hotel booking decision. "
                       "Person 1 has booked a hotel stay arriving on {date} (week {week}), "
                       "with {weekend} weekend night(s) and {weekday} weekday night(s). "
                       "The party consists of {adults} adult(s)[, {children} child(ren)]. "
                       "Person 1 is {repeated_status} and has {prev_cancel} previous cancellation(s). "
                       "They have requested {parking} car parking space(s) and made {special} special request(s).",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (keep booking) or 'PREDICTION: 0' (cancel)."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{pred_response}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is correct if it equals Person 1's true decision."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (keep booking) or 'PREDICTION: 0' (cancel)."),
            ("Critique",
             "SCENARIO:\n{base_prompt}\n\nPROPOSED LOGIC:\n{thought}\n\n"
             "TASK: Critique this reasoning in 1 sentence."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{thought}\n\nCRITIQUE:\n{critique}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is correct if it equals Person 1's true decision."),
        ],
    }),
    ("Uber", {
        "base_prompt": "You are helping Person 1 with a ride decision. "
                       "Person 1 has received a ride request going from {Pickup_Location} to {Drop_Location}.",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (accept) or 'PREDICTION: 0' (decline)."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{pred_response}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is correct if it equals Person 1's true decision."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "TASK: Predict Person 1's decision. Explain your reasoning in 1 sentence. "
             "Then conclude with exactly 'PREDICTION: 1' (accept) or 'PREDICTION: 0' (decline)."),
            ("Critique",
             "SCENARIO:\n{base_prompt}\n\nPROPOSED LOGIC:\n{thought}\n\n"
             "TASK: Critique this reasoning in 1 sentence."),
            ("Escalation",
             "SCENARIO:\n{base_prompt}\n\nPREDICTION & REASONING:\n{thought}\n\nCRITIQUE:\n{critique}\n\n"
             "TASK: Respond with a digit ONLY.\n"
             "Output 0 to implement your prediction, or 1 to escalate to Person 1.\n"
             "A decision is correct if it equals Person 1's true decision."),
        ],
    }),
    ("AIME", {
        "base_prompt": "You are helping Person 1 solve the following math problem: {Question}.",
        "base": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "Solve the math problem step by step, then conclude with 'SOLUTION: <integer>'."),
            ("Escalation",
             "PROBLEM: {base_prompt}\n"
             "SOLUTION AND REASONING: {pred_response}\n\n"
             "Implement your solution (0) or escalate to Person 1 (1). Respond with the digit ONLY."),
        ],
        "adversarial": [
            ("Prediction + Reasoning",
             "{base_prompt}\n\n"
             "Solve the math problem step by step, then conclude with 'SOLUTION: <integer>'."),
            ("Critique",
             "PROBLEM: {base_prompt}\n\n"
             "PROPOSED SOLUTION: {thought}\n\n"
             "Critique this reasoning in 1 sentence."),
            ("Escalation",
             "PROBLEM: {base_prompt}\n"
             "SOLUTION AND REASONING: {thought}\n"
             "CRITIQUE: {critique}\n\n"
             "Implement your solution (0) or escalate to Person 1 (1). Respond with the digit ONLY."),
        ],
    }),
]

for dataset_name, config in datasets:
    doc.add_heading(dataset_name, level=1)

    # Show base_prompt template
    p = doc.add_paragraph()
    run = p.add_run("Scenario template")
    run.bold = True
    run.font.size = Pt(11)
    prompt_para = doc.add_paragraph()
    run = prompt_para.add_run(config["base_prompt"])
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for method in ["base", "adversarial"]:
        doc.add_heading(f"{method.title()} Path", level=2)
        steps = config[method]
        for i, (step_name, prompt_text) in enumerate(steps, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Step {i}: {step_name}")
            run.bold = True
            run.font.size = Pt(11)

            prompt_para = doc.add_paragraph()
            prompt_para.style = doc.styles['Normal']
            run = prompt_para.add_run(prompt_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph()  # spacing

output_path = "results/prompt_reference.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
